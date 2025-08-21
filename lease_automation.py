import os
import json
from docx import Document 
from io import BytesIO
import traceback


def load_sig_block_template(filename):
    """
    Load a signature block template from the templates/sigBlocks directory.
    
    Args:
        filename (str): Name of the template file
        
    Returns:
        str: Template content
    """  
    try:
        # Get the directory where this script is located
        script_dir = os.path.dirname(os.path.abspath(__file__))
        # Use repository root (this file's directory) for current structure
        project_root = script_dir
        path = os.path.join(project_root, 'templates', 'sigBlocks', filename)
        
        with open(path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except FileNotFoundError:
        return f"Template file '{filename}' not found at {path}"
    except Exception as e:
        return f"Error reading template: {str(e)}"
def load_notary_template():
    """
    Load the notary template from templates/Notorary/notrary.txt.
    
    Returns:
        str: Notary template content
    """
    try:
        # Get the directory where this script is located
        script_dir = os.path.dirname(os.path.abspath(__file__))
        # Use repository root (this file's directory) for current structure
        project_root = script_dir
        notary_file_path = os.path.join(project_root, 'templates', 'Notorary', 'notrary.txt')
        
        with open(notary_file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except FileNotFoundError:
        return "Notary block template file 'notrary.txt' not found."
    except Exception as e:
        return f"Error reading notary block: {str(e)}"
def get_sig_block(owner_type, num_signatures):
    """
    Get signature block content based on owner type and number of signatures.
    
    Args:
        owner_type (str): Type of owner (individual, corporation, etc.)
        num_signatures (int): Number of signatures needed
        
    Returns:
        list: Array with filename1 and filename2 content
    """
    filename1 = None
    filename2 = None
    filename1Content = None
    filename2Content = None
    
    # Map owner types to template files (using files that actually exist)
    if owner_type == 'his/her sole property' and num_signatures == 1:
        filename1 = 'SI1.txt'
    elif owner_type == 'a married couple' and num_signatures == 2:
        filename1 = 'I1.txt'
        filename2 = 'I1.txt'
    elif owner_type == 'Corporation':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'LLC':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'LP':    
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'Trust':
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    elif owner_type == 'Sole Owner, married couple' and num_signatures == 2:
        filename1 = 'I1.txt'
        filename2 = 'SI1.txt'
    elif 'individual' in owner_type.lower():
        # Default individual case
        filename1 = 'I1.txt'
        if num_signatures == 2:
            filename2 = 'I1.txt'
    else:
        # Default entity case
        filename1 = 'E1.txt'
        if num_signatures == 2:
            filename2 = 'E1.txt'
    
    # Load content from filename1 if it exists
    if filename1:
        filename1Content = load_sig_block_template(filename1)
    
    # Load content from filename2 if it exists  
    if filename2:
        filename2Content = load_sig_block_template(filename2)
    
    # Return array with filename1 and filename2 content
    return [filename1Content, filename2Content]
def generator(owner_type, is_notary, notary_block, num_signatures):
    """
    Generate signature blocks using the existing generator logic.
    This is the main function used in the /generate_signature_block route.
    
    Args:
        owner_type (str): Type of owner
        is_notary (bool): Whether to include notary
        notary_block (str): Notary block content
        num_signatures (int): Number of signatures
        
    Returns:
        str: Generated signature block(s)
    """
    # Generate notary block with default parameters
    notary_content = load_notary_template()
    
    # Get signature block content
    filecontent = get_sig_block(owner_type, num_signatures)
    sigB1 = filecontent[0] if filecontent[0] is not None else ''
    sigB2 = filecontent[1] if filecontent[1] is not None else ''
    
    final_string = ""

    # Edge case where there are unique sig blocks 
    if owner_type == 'Sole owner, married couple' and is_notary:
        final_string += sigB1
        if is_notary and notary_content:
            final_string += "\n\n" + notary_content
        final_string += "\n\n" + sigB2
        if is_notary and notary_content:
            final_string += "\n\n" + notary_content
        return final_string
    
    if owner_type == 'Sole owner, married couple' and not is_notary:
        final_string += sigB1
        final_string += "\n\n" + sigB2
        return final_string
    
    # Generate additional signature blocks based on num_signatures
    for i in range(num_signatures):
        if is_notary and notary_content:
            final_string += "\n\n" + sigB1
            final_string += "\n\n" + notary_content
        else:
            final_string += "\n\n" + sigB1

    return final_string
def generate_signature_blocks_from_json(json_data):
    """
    Generate signature blocks with and without notary from JSON data.
    Uses the existing generator() function twice for proper dynamic/nested functionality.
    
    Args:
        json_data (dict): JSON data containing grantor information
        
    Returns:
        dict: JSON response with both signature blocks
    """
    try:
        # Step 1: Extract information from JSON using getter methods
        grantor_name = json_data.get('grantor_name', '')
        grantor_name_1 = json_data.get('grantor_name_1', '')
        grantor_name_2 = json_data.get('grantor_name_2', '')
        trust_entity_name = json_data.get('trust_entity_name', '')
        owner_type = json_data.get('owner_type', 'individual')
        number_of_signatures = json_data.get('number_of_grantor_signatures', 1)
        state = json_data.get('state', '')
        county = json_data.get('county', '')
        
        # Step 2: Call the generator() function twice
        # First call: WITHOUT notary (isNotary = False)
        signature_block_without_notary = generator(
            owner_type=owner_type,
            is_notary=False,  # isNotary = False
            notary_block='',
            num_signatures=number_of_signatures
        )
        
        # Second call: WITH notary (isNotary = True)
        signature_block_with_notary = generator(
            owner_type=owner_type,
            is_notary=True,   # isNotary = True
            notary_block='',
            num_signatures=number_of_signatures
        )
        
        # Step 3: Return both generated blocks in JSON format (standardized keys)
        return {
            "Signature_block": signature_block_without_notary,
            "Signature_Block_With_Notary": signature_block_with_notary
        }
        
    except Exception as e:
        # Return error in the same format
        return {
            "signature_block": f"Error generating signature block: {str(e)}",
            "Signature Block With Notrary": f"Error generating signature block with notary: {str(e)}"
        }
def build_exhibit_string_from_json(json_data):
   
    try:
        print("[DEBUG] Starting exhibit string generation from JSON")
        
        # Parse JSON if it's a string
        if isinstance(json_data, str):
            import json
            data = json.loads(json_data)
        else:
            data = json_data
        
        # Extract document information
        document_name = data.get("document_name", "Unknown Document")
        grantor_type = data.get("grantor_type", "Unknown")
        grantor_name = data.get("grantor_name", "Unknown")
        state = data.get("state", "Unknown")
        county = data.get("county", "Unknown")
        total_acres = data.get("total_acres", 0)
        number_of_parcels = data.get("number_of_parcels", 0)
        
        # Extract and create parcel objects
        raw_parcels = data.get("parcels", [])
        
        print(f"[DEBUG] Processing document: {document_name}")
        print(f"[DEBUG] Found {len(raw_parcels)} parcels, total acres: {total_acres}")
        
        # Validate parcels data
        if not isinstance(raw_parcels, list) or len(raw_parcels) == 0:
            raise ValueError("Parcels must be a non-empty list")
        
        # Create parcel objects from the JSON data
        parcel_objects = []
        for i, raw_parcel in enumerate(raw_parcels, 1):
            if not isinstance(raw_parcel, dict):
                print(f"[WARNING] Invalid parcel data at index {i}: {raw_parcel}")
                continue
            
            # Create parcel object with all the data
            parcel_obj = {
                "parcelNumber": i,
                "apn": raw_parcel.get("apn", f"Unknown-{i}"),
                "acres": raw_parcel.get("acres", 0),
                "legal_description": raw_parcel.get("legal_description", "No legal description provided"),
                "isPortion": raw_parcel.get("isPortion", False),
                "templateType": "standard"  # Default template type
            }
            
            parcel_objects.append(parcel_obj)
            print(f"[DEBUG] Created parcel object {i}: APN {parcel_obj['apn']}, {parcel_obj['acres']} acres, isPortion: {parcel_obj['isPortion']}")
        
        # Now use the existing build_exhibit_string function with our parcel objects
        exhibit_string = build_exhibit_string(parcel_objects)
        
        # Create the output JSON structure
        output_json = {
            "document_name": document_name,
            "grantor_type": grantor_type,
            "grantor_name": grantor_name,
            "state": state,
            "county": county,
            "total_acres": total_acres,
            "number_of_parcels": number_of_parcels,
            "exhibit_a_string": exhibit_string,
            "parcels_processed": len(parcel_objects),
            "parcel_objects": parcel_objects,  # Include the created parcel objects
            "generation_timestamp": __import__('datetime').datetime.now().isoformat()
        }
        
        print(f"[DEBUG] Generated exhibit string, length: {len(exhibit_string)}")
        print(f"[DEBUG] Output JSON created with {len(output_json)} fields")
        print(f"[DEBUG] Created {len(parcel_objects)} parcel objects")
        
        return output_json
        
    except Exception as e:
        print(f"[ERROR] Failed to build exhibit string from JSON: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Return error JSON
        error_json = {
            "error": str(e),
            "error_type": type(e).__name__,
            "generation_timestamp": __import__('datetime').datetime.now().isoformat()
        }
        return error_json
def build_exhibit_string(parcels):

    """
    Build the Exhibit A text string from parcel objects.
    
    Args:
        parcels: List of parcel objects with parcelNumber, isPortion, and templateType properties
    
    Returns:
        str: The complete Exhibit A text string
    """
    try:
        print(f"[DEBUG] Building exhibit string for {len(parcels)} parcels")
        
        # Validate parcels data
        if not isinstance(parcels, list) or len(parcels) == 0:
            raise ValueError("Parcels must be a non-empty list")
        
        # Start with header
        exhibit_parts = ["EXHIBIT A", "", "General Description of Property", ""]
        
        # Add image placeholder
        exhibit_parts.append("[Image]")
        exhibit_parts.append("")
        
        # Add parcel descriptions using the parcel objects
        for parcel in parcels:
            if not isinstance(parcel, dict):
                print(f"[WARNING] Invalid parcel object: {parcel}")
                continue
            
            parcel_number = parcel.get("parcelNumber", "Unknown")
            apn = parcel.get("apn", "Unknown")
            acres = parcel.get("acres", 0)
            legal_description = parcel.get("legal_description", "No legal description provided")
            is_portion = parcel.get("isPortion", False)
            
            # Create custom description based on whether it's a portion or parcel
            if is_portion:
                parcel_description = f"Portion {parcel_number} (APN: {apn}, {acres} acres):\n\n{legal_description}"
            else:
                parcel_description = f"Parcel {parcel_number} (APN: {apn}, {acres} acres):\n\n{legal_description}"
            
            print(f"[DEBUG] Processing {parcel_number}: APN {apn}, {acres} acres, isPortion: {is_portion}")
            exhibit_parts.append(parcel_description)
            exhibit_parts.append("")  # Add spacing between parcels
        
        # Join all parts
        exhibit_string = "\n".join(exhibit_parts)
        
        print(f"[DEBUG] Generated exhibit string, length: {len(exhibit_string)}")
        return exhibit_string
        
    except Exception as e:
        print(f"[ERROR] Failed to build exhibit string: {str(e)}")
        import traceback
        traceback.print_exc()
        raise

def update_json_with_generated_content(json_data):
    """
    Update the JSON data with generated signature blocks and exhibit A content.
    
    Args:
        json_data (dict): The JSON data dictionary to update
        
    Returns:
        dict: Updated JSON data with signature blocks and exhibit A
    """
    # Generate signature blocks and exhibit A content
    sigBlocks = generate_signature_blocks_from_json(json_data)
    exhibitA = build_exhibit_string_from_json(json_data)
    
    # Add the two signature blocks individually (use standardized keys)
    json_data["Signature_block"] = sigBlocks.get("Signature_block")
    # Backward compatibility for any older key names
    json_data["Signature_Block_With_Notary"] = sigBlocks.get("Signature_Block_With_Notary") or sigBlocks.get("Signature Block With Notrary") or sigBlocks.get("Signature Block With Notary")
    
    # Add the exhibit A content
    json_data["exhibit_a"] = exhibitA
    
    return json_data
def keyValueMapping(json_data):
    """
    Create a key-value mapping from the JSON data.
    
    Args:
        json_data (dict): The JSON data dictionary to process
        
    Returns:
        list: List of dictionaries with 'key' and 'value' pairs for document replacement
    """
    mapping_data = []
    
    # Basic field mappings
    basic_fields = {
        "document_name": "[Document Name]",
        "grantor_type": "[Grantor Type]", 
        "grantor_name_1": "[Grantor Name 1]",
        "grantor_name_2": "[Grantor Name 2]",
        "trust_entity_name": "[Trust Entity Name]",
        "grantor_name": "[Grantor Name]",
        "owner_type": "[Owner Type]",
        "number_of_grantor_signatures": "[Number of Grantor Signatures]",
        "grantor_address_1": "[Grantor Address 1]",
        "grantor_address_2": "[Grantor Address 2]",
        "state": "[State]",
        "county": "[County]",
        "total_acres": "[Total Acres]",
        "number_of_parcels": "[Number of Parcels]"
    }
    
    # Add basic field mappings
    for json_key, placeholder in basic_fields.items():
        if json_key in json_data:
            value = json_data[json_key]
            if value is not None:
                # Convert to string and check if it's not empty
                str_value = str(value).strip()
                if str_value != "" and str_value.lower() != "na":
                    mapping_data.append({
                        "key": placeholder,
                        "value": str_value
                    })
    
    # Add signature blocks (use existing values from JSON if available)
    if "Signature_Block_With_Notary" in json_data:
        mapping_data.append({
            "key": "[Signature Block With Notary]",
            "value": json_data["Signature_Block_With_Notary"]
        })
    else:
        mapping_data.append({
            "key": "[Signature Block With Notary]",
            "value": ""
        })
    
    if "Signature_block" in json_data:
        mapping_data.append({
            "key": "[Signature Block]",
            "value": json_data["Signature_block"]
        })
    else:
        mapping_data.append({
            "key": "[Signature Block]",
            "value": ""
        })
    
    # Add APN list
    if "apn_list" in json_data and json_data["apn_list"]:
        mapping_data.append({
            "key": "[APN List]",
            "value": json_data["apn_list"]
        })
    
    # Add Exhibit A mappings
    if "exhibit_a" in json_data:
        exhibit = json_data["exhibit_a"]
        
        # Basic exhibit fields
        exhibit_basic_fields = {
            "county": "[Exhibit A - County]",
            "document_name": "[Exhibit A - Document Name]",
            "exhibit_a_string": "[Exhibit A - Exhibit A String]",
            "generation_timestamp": "[Exhibit A - Generation Timestamp]",
            "grantor_name": "[Exhibit A - Grantor Name]",
            "grantor_type": "[Exhibit A - Grantor Type]",
            "number_of_parcels": "[Exhibit A - Number of Parcels]",
            "state": "[Exhibit A - State]",
            "total_acres": "[Exhibit A - Total Acres]"
        }
        
        for json_key, placeholder in exhibit_basic_fields.items():
            if json_key in exhibit:
                value = exhibit[json_key]
                if value is not None:
                    mapping_data.append({
                        "key": placeholder,
                        "value": value
                    })
        
        # Add parcel-specific mappings from exhibit_a.parcel_objects
        if "parcel_objects" in exhibit and exhibit["parcel_objects"]:
            for i, parcel in enumerate(exhibit["parcel_objects"], 1):
                parcel_prefix = f"[Exhibit A - Parcel {i}"
                mapping_data.extend([
                    {"key": f"{parcel_prefix} APN]", "value": parcel.get("apn", "")},
                    {"key": f"{parcel_prefix} Acres]", "value": parcel.get("acres", 0)},
                    {"key": f"{parcel_prefix} Is Portion]", "value": parcel.get("isPortion", False)},
                    {"key": f"{parcel_prefix} Legal Description]", "value": parcel.get("legal_description", "")},
                    {"key": f"{parcel_prefix} Parcel Number]", "value": parcel.get("parcelNumber", i)},
                    {"key": f"{parcel_prefix} Template Type]", "value": parcel.get("templateType", "standard")}
                ])
    
    # Add Parcels section mappings (from main parcels array)
    if "parcels" in json_data and json_data["parcels"]:
        for i, parcel in enumerate(json_data["parcels"], 1):
            parcel_prefix = f"[Parcels - Parcel {i}"
            mapping_data.extend([
                {"key": f"{parcel_prefix} APN]", "value": parcel.get("apn", "")},
                {"key": f"{parcel_prefix} Acres]", "value": parcel.get("acres", 0)},
                {"key": f"{parcel_prefix} Is Portion]", "value": parcel.get("isPortion", False)},
                {"key": f"{parcel_prefix} Legal Description]", "value": parcel.get("legal_description", "")},
                {"key": f"{parcel_prefix} Parcel Number]", "value": parcel.get("parcelNumber", i)}
            ])
    
    return mapping_data
def getMapping(json_data):
    mapping = keyValueMapping(update_json_with_generated_content(json_data))
    return mapping


def replace_placeholders_in_document(doc, mapping, track_changes=False):
    """
    Core function that replaces placeholders in a DOCX document.
    Supports both normal replacement and track changes mode.
    
    Args:
        doc: Document object to process
        mapping: Dictionary of placeholder keys to replacement values
        track_changes: If True, adds "NEW:" prefix and highlights changes
    
    Returns:
        Document: Processed document with replacements
    """
    try:
        print(f"[DEBUG] Starting placeholder replacement. Track changes: {track_changes}")
        print(f"[DEBUG] Processing {len(mapping)} placeholders")
        
        # Apply track changes prefix if enabled
        if track_changes:
            mapping = {k: f"NEW:{v}" for k, v in mapping.items()}
            print("[DEBUG] Applied 'NEW:' prefix for track changes")
        
        # Choose replacement method based on track changes setting
        if track_changes:
            doc = _replace_placeholders_with_track_changes(doc, mapping)
        else:
            doc = _replace_placeholders_normal(doc, mapping)
        
        print("[DEBUG] Placeholder replacement completed successfully")
        return doc
        
    except Exception as e:
        print(f"[ERROR] Failed to replace placeholders: {str(e)}")
        import traceback
        traceback.print_exc()
        raise


def _replace_placeholders_normal(doc, mapping):
    """
    Normal placeholder replacement without track changes.
    Replaces placeholders throughout the entire document.
    
    Args:
        doc: Document object to process
        mapping: Dictionary of placeholder keys to replacement values
    
    Returns:
        Document: Processed document with replacements
    """
    def replace_in_runs(runs, mapping):
        """Replace placeholders in a sequence of runs"""
        if not runs:
            return
            
        # Join all run text together
        full_text = ''.join(run.text for run in runs)
        
        # Apply all replacements
        for key, value in mapping.items():
            if not value.strip():
                continue
            full_text = full_text.replace(key, value)
        
        # Update the runs with the replaced text
        if runs:
            runs[0].text = full_text
            # Clear all other runs
            for run in runs[1:]:
                run.text = ''
    
    def process_paragraph(paragraph, mapping):
        """Process a single paragraph for placeholders"""
        if not paragraph.runs:
            return
            
        # Check if this paragraph contains any placeholders
        joined = ''.join(run.text for run in paragraph.runs)
        if any(key in joined for key in mapping.keys()):
            replace_in_runs(paragraph.runs, mapping)
    
    def process_table(table, mapping):
        """Process a table for placeholders"""
        for row in table.rows:
            for cell in row.cells:
                process_block(cell, mapping)
    
    def process_block(block, mapping):
        """Process a block (paragraphs and tables) for placeholders"""
        for paragraph in block.paragraphs:
            process_paragraph(paragraph, mapping)
        for table in getattr(block, 'tables', []):
            process_table(table, mapping)
    
    # Process all document sections
    print(f"[DEBUG] Processing {len(doc.paragraphs)} paragraphs")
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, mapping)
    
    print(f"[DEBUG] Processing {len(doc.tables)} tables")
    for table in doc.tables:
        process_table(table, mapping)
    
    # Process headers and footers
    for section in doc.sections:
        if hasattr(section, 'header'):
            process_block(section.header, mapping)
        if hasattr(section, 'footer'):
            process_block(section.footer, mapping)
    
    # Process footnotes if they exist
    if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
        try:
            for footnote in doc.part.footnotes.part.footnotes:
                for paragraph in footnote.paragraphs:
                    process_paragraph(paragraph, mapping)
        except Exception as e:
            print(f"[WARNING] Could not process footnotes: {str(e)}")
    
    return doc


def _replace_placeholders_with_track_changes(doc, mapping):
    """
    Placeholder replacement with track changes highlighting.
    Adds yellow highlighting to all replaced text.
    
    Args:
        doc: Document object to process
        mapping: Dictionary of placeholder keys to replacement values
    
    Returns:
        Document: Processed document with highlighted replacements
    """
    def process_paragraph(paragraph, mapping):
        """Process a single paragraph for placeholders with highlighting"""
        for run in paragraph.runs:
            for key, value in mapping.items():
                if not value.strip():
                    continue
                    
                # Check if this run contains the placeholder
                if key in run.text:
                    # Replace the placeholder with the value
                    run.text = run.text.replace(key, value)
                    # Highlight the replacement with yellow
                    run.font.highlight_color = 7  # 7 = yellow highlight
                    print(f"[DEBUG] Highlighted replacement: {key} -> {value[:30]}{'...' if len(value) > 30 else ''}")
                    break  # Only process one replacement per run to avoid conflicts
    
    def process_table(table, mapping):
        """Process a table for placeholders with highlighting"""
        for row in table.rows:
            for cell in row.cells:
                process_block(cell, mapping)
    
    def process_block(block, mapping):
        """Process a block (paragraphs and tables) for placeholders with highlighting"""
        for paragraph in block.paragraphs:
            process_paragraph(paragraph, mapping)
        for table in getattr(block, 'tables', []):
            process_table(table, mapping)
    
    # Process all document sections
    print(f"[DEBUG] Processing {len(doc.paragraphs)} paragraphs with track changes")
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, mapping)
    
    print(f"[DEBUG] Processing {len(doc.tables)} tables with track changes")
    for table in doc.tables:
        process_table(table, mapping)
    
    # Process headers and footers
    for section in doc.sections:
        if hasattr(section, 'header'):
            process_block(section.header, mapping)
        if hasattr(section, 'footer'):
            process_block(section.footer, mapping)
    
    # Process footnotes if they exist
    if hasattr(doc, 'part') and hasattr(doc.part, 'footnotes'):
        try:
            for footnote in doc.part.footnotes.part.footnotes:
                for paragraph in footnote.paragraphs:
                    process_paragraph(paragraph, mapping)
        except Exception as e:
            print(f"[WARNING] Could not process footnotes with track changes: {str(e)}")
    
    return doc


def simple_document_replacement(docx_file, mapping_json, output_filename='processed_document.docx', track_changes=False):
    """
    Simple document replacement function that takes JSON mapping and DOCX template,
    performs text replacement (with optional track changes), and returns the processed DOCX file.
    
    Args:
        docx_file: File path or file-like object to the DOCX template
        mapping_json: JSON string in the format [{"key": "...", "value": "..."}]
        output_filename: Name for the output file (optional)
        track_changes: If True, enables track changes mode with highlighting
    
    Returns:
        tuple: (success: bool, result: str or bytes, error_message: str)
        - If success=True: result contains the DOCX file bytes
        - If success=False: result is None, error_message contains the error
    """
    try:
        import json
        from docx import Document
        from io import BytesIO
        
        print(f"[DEBUG] Starting document replacement for: {output_filename}")
        print(f"[DEBUG] Track changes enabled: {track_changes}")
        
        # Parse the JSON mapping
        if isinstance(mapping_json, str):
            mapping_data = json.loads(mapping_json)
        else:
            mapping_data = mapping_json
        
        # Validate mapping format
        if not isinstance(mapping_data, list):
            return False, None, "Mapping must be a list of key-value objects"
        
        # Convert to the format expected by replacement functions
        mapping = {}
        for item in mapping_data:
            if isinstance(item, dict) and 'key' in item and 'value' in item:
                key = item['key'].strip()
                value = str(item['value']).strip()
                if key and value:  # Only add non-empty key-value pairs
                    mapping[key] = value
                    print(f"[DEBUG] Added mapping: {key} -> {value[:50]}{'...' if len(value) > 50 else ''}")
        
        if not mapping:
            return False, None, "No valid key-value pairs found in mapping"
        
        print(f"[DEBUG] Processed {len(mapping)} key-value pairs")
        
        # Load the DOCX document
        if hasattr(docx_file, 'read'):
            # File-like object
            doc = Document(docx_file)
        else:
            # File path
            doc = Document(docx_file)
        
        print(f"[DEBUG] Loaded DOCX document with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        # Perform placeholder replacement
        doc = replace_placeholders_in_document(doc, mapping, track_changes)
        
        # Save the processed document to bytes
        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        
        docx_bytes = output_stream.getvalue()
        print(f"[DEBUG] Document processed successfully. Output size: {len(docx_bytes)} bytes")
        
        return True, docx_bytes, ""
        
    except json.JSONDecodeError as e:
        error_msg = f"Invalid JSON format: {str(e)}"
        print(f"[ERROR] {error_msg}")
        return False, None, error_msg
    except Exception as e:
        error_msg = f"Document processing failed: {str(e)}"
        print(f"[ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return False, None, error_msg


